from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment


def ensure_export_dir(export_dir):
    path = Path(export_dir)
    path.mkdir(parents=True, exist_ok=True)
    return path


def export_forecasts_to_xlsx(export_dir, forecasts):
    export_path = ensure_export_dir(export_dir) / f"forecast_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Прогноз цен"
    headers = [
        "Год", "Организация", "Направление", "Модель", "Прогнозная цена",
        "Рекомендация"
    ]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center")
    for forecast in forecasts:
        ws.append([
            forecast.year,
            forecast.organization,
            forecast.program_name,
            forecast.model_name,
            forecast.predicted_price,
            forecast.recommendation,
        ])
    for column_cells in ws.columns:
        max_length = max(len(str(cell.value)) if cell.value is not None else 0 for cell in column_cells)
        ws.column_dimensions[column_cells[0].column_letter].width = min(max_length + 2, 60)
    wb.save(export_path)
    return export_path


def export_forecasts_to_docx(export_dir, forecasts):
    export_path = ensure_export_dir(export_dir) / f"forecast_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    document.add_heading("Отчет по прогнозированию ценообразования на образовательные услуги", level=1)
    document.add_paragraph(
        "Документ сформирован автоматически на основе результатов работы модуля прогнозирования. "
        "Прогнозная цена рассчитана с учетом характеристик образовательной программы, формата обучения, "
        "региона, спроса, конкурентной цены и дополнительных экономических показателей."
    )
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Год", "Организация", "Направление", "Модель", "Цена", "Рекомендация"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for forecast in forecasts:
        row = table.add_row().cells
        row[0].text = str(forecast.year)
        row[1].text = forecast.organization
        row[2].text = forecast.program_name
        row[3].text = forecast.model_name
        row[4].text = f"{forecast.predicted_price:,.2f}".replace(",", " ")
        row[5].text = forecast.recommendation or ""
    document.add_paragraph(
        "Полученные значения следует использовать как аналитический ориентир, а не как окончательное управленческое решение. "
        "Перед утверждением цены необходимо учитывать ограничения бюджета, план набора, позиционирование программы и текущую конкурентную среду."
    )
    document.save(export_path)
    return export_path


def export_dataset_summary_to_xlsx(export_dir, summary, dynamics, rating):
    export_path = ensure_export_dir(export_dir) / f"dataset_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    with pd.ExcelWriter(export_path, engine="openpyxl") as writer:
        pd.DataFrame([summary]).to_excel(writer, sheet_name="Сводка", index=False)
        pd.DataFrame(dynamics).to_excel(writer, sheet_name="Динамика", index=False)
        pd.DataFrame(rating).to_excel(writer, sheet_name="Рейтинг программ", index=False)
    return export_path


def export_dataset_summary_to_docx(export_dir, summary, dynamics, rating, formats=None, competitors=None):
    export_path = ensure_export_dir(export_dir) / f"price_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    document.add_heading("Отчет по анализу цен на образовательные услуги", level=1)
    document.add_paragraph(
        "Отчет содержит ключевые показатели текущих цен, динамику по годам, сравнение с конкурентами "
        "и рейтинг программ по спросу."
    )
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Показатель"
    table.rows[0].cells[1].text = "Значение"
    for key, value in summary.items():
        if isinstance(value, list):
            value = ", ".join(map(str, value))
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)

    document.add_heading("Динамика цен", level=2)
    dyn_table = document.add_table(rows=1, cols=5)
    dyn_table.style = "Table Grid"
    headers = ["Год", "Средняя цена", "Рост, %", "Средний прием", "Цена других вузов"]
    for i, header in enumerate(headers):
        dyn_table.rows[0].cells[i].text = header
    for item in dynamics:
        row = dyn_table.add_row().cells
        row[0].text = str(item.get("year", ""))
        row[1].text = str(item.get("mean_price", ""))
        row[2].text = str(item.get("growth_percent", ""))
        row[3].text = str(item.get("mean_students", ""))
        row[4].text = str(item.get("mean_competitor_price", ""))

    document.add_heading("Направления с высоким платным приемом", level=2)
    rank_table = document.add_table(rows=1, cols=4)
    rank_table.style = "Table Grid"
    for i, header in enumerate(["Направление", "Средняя цена", "Средний балл", "Зачислено"]):
        rank_table.rows[0].cells[i].text = header
    for item in rating[:10]:
        row = rank_table.add_row().cells
        row[0].text = str(item.get("program_name", ""))
        row[1].text = str(item.get("mean_price", ""))
        row[2].text = str(item.get("mean_admission_score", ""))
        row[3].text = str(item.get("students", ""))

    document.add_paragraph(
        "Вывод: итоговая цена должна рассматриваться вместе с уровнем спроса, длительностью обучения, "
        "форматом программы, скидками и средними ценами конкурентов."
    )
    document.save(export_path)
    return export_path


def build_report_context(forecasts):
    prices = [forecast.predicted_price for forecast in forecasts]
    if not prices:
        return {"count": 0, "mean": 0, "min": 0, "max": 0}
    return {
        "count": len(prices),
        "mean": round(sum(prices) / len(prices), 2),
        "min": round(min(prices), 2),
        "max": round(max(prices), 2),
    }
